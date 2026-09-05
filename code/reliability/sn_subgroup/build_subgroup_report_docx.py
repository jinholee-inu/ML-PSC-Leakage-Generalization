#!/usr/bin/env python3
"""Build the visually verified Word report for the Sn subgroup analysis."""

from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
RESULTS = ROOT / "outputs" / "psc-subgroup-calibration-moe"
OUTPUT = RESULTS / "PSC_Sn_Subgroup_Calibration_MoE_Report.docx"
FIGURE = RESULTS / "Figure9_subgroup_calibration_moe.png"

NAVY = "18324B"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "20252B"
MUTED = "5E6872"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE_GOLD = "FFF4D6"
GOLD = "7A5A00"
WHITE = "FFFFFF"
RED = "9B1C1C"
GREEN = "176B4D"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(
    run,
    name: str = "Calibri",
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    settings = {
        "Title": (24, NAVY, 0, 5),
        "Subtitle": (12.5, MUTED, 0, 12),
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in settings.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.line_spacing = 1.0

    for list_name in ["List Bullet", "List Number"]:
        style = doc.styles[list_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167


def set_page_layout(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)


def add_page_field(paragraph) -> None:
    paragraph.add_run("Page ")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def configure_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("PSC MODEL RELIABILITY  |  SUBGROUP ANALYSIS")
    set_run_font(run, size=8.5, color=MUTED, bold=True)

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_before = Pt(0)
    add_page_field(paragraph)
    for run in paragraph.runs:
        set_run_font(run, size=8.5, color=MUTED)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    if sum(widths) != TABLE_WIDTH_DXA:
        raise ValueError(f"Table widths sum to {sum(widths)}, expected {TABLE_WIDTH_DXA}")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def style_table(table, widths: list[int], numeric_columns: set[int] | None = None) -> None:
    numeric_columns = numeric_columns or set()
    set_table_geometry(table, widths)
    table.style = "Table Grid"
    set_repeat_table_header(table.rows[0])
    for row_index, row in enumerate(table.rows):
        prevent_row_split(row)
        if row_index == 0:
            for cell in row.cells:
                set_cell_shading(cell, LIGHT_GRAY)
        for column_index, cell in enumerate(row.cells):
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.05
                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER if column_index in numeric_columns else WD_ALIGN_PARAGRAPH.LEFT
                )
                for run in paragraph.runs:
                    set_run_font(
                        run,
                        size=8.6 if row_index else 8.8,
                        color=INK,
                        bold=row_index == 0,
                    )


def add_table_caption(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    set_run_font(run, size=9.5, color=DARK_BLUE, bold=True)


def add_note(doc: Document, label: str, text: str, fill: str = PALE_GOLD) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(9)
    paragraph.paragraph_format.left_indent = Inches(0.16)
    paragraph.paragraph_format.right_indent = Inches(0.16)
    paragraph.paragraph_format.keep_together = True
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    p_pr.append(shading)
    border = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), GOLD if fill == PALE_GOLD else BLUE)
    border.append(left)
    p_pr.append(border)
    lead = paragraph.add_run(f"{label}  ")
    set_run_font(lead, size=10.5, color=GOLD if fill == PALE_GOLD else DARK_BLUE, bold=True)
    body = paragraph.add_run(text)
    set_run_font(body, size=10.5, color=INK)


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.add_run(text)


def add_body(doc: Document, text: str, bold_lead: str | None = None) -> None:
    paragraph = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        set_run_font(lead, bold=True, color=INK)
        paragraph.add_run(text[len(bold_lead):])
    else:
        paragraph.add_run(text)


def set_image_alt_text(paragraph, description: str) -> None:
    for doc_pr in paragraph._p.xpath(".//wp:docPr"):
        doc_pr.set("descr", description)
        doc_pr.set("title", "Figure 9")


def format_ci(value: float, low: float, high: float, digits: int = 3) -> str:
    return f"{value:.{digits}f} [{low:.{digits}f}, {high:.{digits}f}]"


def build_document() -> Document:
    metrics = pd.read_csv(RESULTS / "subgroup_future_metrics.csv")
    paired = pd.read_csv(RESULTS / "subgroup_paired_comparisons.csv")
    support = pd.read_csv(RESULTS / "subgroup_support.csv")
    upper = pd.read_csv(RESULTS / "subgroup_PCE_upper_tail.csv")
    robustness = pd.read_csv(RESULTS / "subgroup_robustness_strata.csv")
    manifest = json.loads((RESULTS / "subgroup_calibration_moe_run_manifest.json").read_text(encoding="utf-8"))
    verification = json.loads((RESULTS / "independent_subgroup_calibration_moe_verification.json").read_text(encoding="utf-8"))

    doc = Document()
    configure_styles(doc)
    set_page_layout(doc)
    configure_header_footer(doc)
    doc.core_properties.title = "Sn-only and mixed Pb-Sn subgroup calibration and mixture-of-experts analysis"
    doc.core_properties.subject = "Leakage-controlled chronological validation of PSC subgroup models"
    doc.core_properties.author = "PSC reliability analysis workflow"
    doc.core_properties.keywords = "perovskite solar cells, Sn, Pb-Sn, calibration, mixture of experts, DOI bootstrap"

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(2)
    run = kicker.add_run("TECHNICAL FINDINGS REPORT  |  29 AUGUST 2026")
    set_run_font(run, size=9.5, color=BLUE, bold=True)
    title = doc.add_paragraph(style="Title")
    title.add_run("Subgroup Calibration and Mixture-of-Experts for Sn-Containing PSCs")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("Sn-only and mixed Pb-Sn domains under publication-disjoint development and an untouched 2019-2021 test")

    add_note(
        doc,
        "DEPLOYMENT DECISION",
        "Retain the frozen full-DOI-balanced global Random Forest for both Sn-only and mixed Pb-Sn devices. "
        "The one-standard-error rule selected identity calibration and a zero expert mixture weight for every target. "
        "Subgroup experts significantly worsened future PCE MAE rather than repairing the observed temporal bias.",
        fill=LIGHT_BLUE,
    )

    doc.add_heading("Executive result", level=1)
    add_body(
        doc,
        "This analysis asked whether the poor transfer observed in Sn-containing absorber domains could be repaired by post-hoc subgroup calibration or by a domain-specific Random Forest blended with the frozen global model. All tuning used only records published through 2018 and DOI-disjoint out-of-fold predictions. The 2019-2021 outcomes were evaluated once after the calibrator, expert configuration, and mixture weight had been fixed.",
    )
    add_bullet(doc, "Sn-only support: 389 historical records from 83 DOI groups; 155 future records from 30 DOI groups.")
    add_bullet(doc, "Mixed Pb-Sn support: 274 historical records from 57 DOI groups; 98 future records from 25 DOI groups.")
    add_bullet(doc, "For all eight domain-target combinations, the one-standard-error selection rule retained identity calibration, alpha = 0, and the frozen global policy.")
    add_bullet(doc, "Independent verification passed 22 of 22 integrity checks, including exact frozen-prediction recovery, zero DOI overlap, unique prediction keys, and metric recomputation to 1.78e-15.")

    doc.add_heading("Study design", level=1)
    add_body(
        doc,
        "The reference model was the previously selected multi-output Random Forest trained with full inverse publication weighting (1/n_DOI). DOI, publication year, and photovoltaic targets were excluded from its feature matrix. Historical model-selection predictions were regenerated using the frozen five DOI folds, but every fold was restricted to publications through 2018. This prevented future publications from entering preprocessing, training, calibration, or expert selection.",
    )
    add_body(
        doc,
        "Three subgroup adaptations were evaluated. The calibrator grid contained identity, bias-only, and ridge-shrunk affine corrections. Four domain-expert Random Forest configurations varied category frequency thresholds and leaf regularization. The mixture used p_mix = (1-alpha)p_global + alpha p_expert with alpha in increments of 0.05. Selection was based on DOI-balanced historical out-of-fold error and a one-standard-error rule that favored stronger shrinkage and lower complexity. Final uncertainty comparisons used 1,000 paired DOI-cluster bootstrap replicates on identical future records.",
    )

    add_table_caption(doc, "Table 1. Subgroup support before and after the chronological cutoff")
    table = doc.add_table(rows=1, cols=4)
    headers = ["Domain", "Period", "Records", "DOI groups"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
    for domain in ["Sn-only", "Mixed Pb-Sn"]:
        for period in ["Historical <=2018", "Future 2019-2021"]:
            row = support.loc[(support["domain"] == domain) & (support["period"] == period)].iloc[0]
            cells = table.add_row().cells
            values = [domain, period, f"{int(row.records):,}", f"{int(row.DOI_groups):,}"]
            for cell, text in zip(cells, values):
                cell.text = text
    style_table(table, [1800, 3400, 1900, 2260], numeric_columns={2, 3})

    doc.add_heading("Figure 9. Model selection and chronological transfer", level=1)
    figure_paragraph = doc.add_paragraph()
    figure_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    figure_paragraph.paragraph_format.keep_with_next = True
    figure_run = figure_paragraph.add_run()
    figure_run.add_picture(str(FIGURE), width=Inches(5.85))
    set_image_alt_text(
        figure_paragraph,
        "Four-panel scientific figure showing subgroup DOI support, historical out-of-fold MAE ratios, future paired MAE changes, and PCE calibration for Sn-only and mixed Pb-Sn perovskite solar cells.",
    )
    caption = doc.add_paragraph()
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.keep_with_next = False
    run = caption.add_run(
        "Figure 9. Leakage-controlled subgroup calibration and shrinkage mixture-of-experts analysis for Sn-containing perovskite solar cells. "
        "(a) Numbers of independent DOI groups in the historical development and future chronological cohorts; labels report device records. "
        "(b) Historical publication-disjoint out-of-fold MAE relative to the frozen global model after the one-standard-error selection rule. "
        "(c) Paired percentage change in future publication-balanced MAE relative to the global model. Points and horizontal bars denote estimates and paired DOI-cluster bootstrap 95% confidence intervals (1,000 replicates). "
        "(d) Publication-balanced PCE calibration in four prediction bins. Because identity calibration and alpha = 0 were selected for all targets, the development-selected curves coincide with the frozen global curves."
    )
    set_run_font(run, size=9.3, color=MUTED, italic=True)

    doc.add_heading("Independent chronological results", level=1)
    add_table_caption(doc, "Table 2. Future publication-balanced PCE performance")
    table = doc.add_table(rows=1, cols=6)
    headers = ["Domain", "Method", "R2 [95% CI]", "MAE [95% CI]", "Bias [95% CI]", "MAE change vs global"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
    for domain in ["Sn-only", "Mixed Pb-Sn"]:
        for method in ["Frozen global", "Domain expert"]:
            row = metrics.loc[
                metrics["domain"].eq(domain)
                & metrics["target"].eq("PCE")
                & metrics["method"].eq(method)
                & metrics["evaluation"].eq("Publication-balanced")
            ].iloc[0]
            if method == "Frozen global":
                change = "Reference"
            else:
                compare = paired.loc[
                    paired["domain"].eq(domain)
                    & paired["target"].eq("PCE")
                    & paired["method"].eq(method)
                    & paired["evaluation"].eq("Publication-balanced")
                ].iloc[0]
                change = f"{compare.MAE_change_percent:+.1f}% [{compare.MAE_change_percent_CI_low:+.1f}, {compare.MAE_change_percent_CI_high:+.1f}]"
            values = [
                domain,
                method,
                format_ci(row.R2, row.R2_CI_low, row.R2_CI_high),
                format_ci(row.MAE, row.MAE_CI_low, row.MAE_CI_high),
                format_ci(row.bias, row.bias_CI_low, row.bias_CI_high),
                change,
            ]
            cells = table.add_row().cells
            for cell, text in zip(cells, values):
                cell.text = text
    style_table(table, [1350, 1560, 1600, 1700, 1650, 1500], numeric_columns={2, 3, 4, 5})

    add_body(
        doc,
        "Sn-only. The frozen global model yielded publication-balanced PCE R2 = 0.230 and MAE = 1.985 percentage points, with a negative bias of -0.997 percentage points. The domain expert reduced R2 by 0.169 and increased MAE by 14.4% (95% CI, 8.0% to 22.1%). Its nearly unchanged mean bias shows that the extra error arose mainly from degraded ranking and local fit, not from successful correction of the average offset.",
        bold_lead="Sn-only.",
    )
    add_body(
        doc,
        "Mixed Pb-Sn. The frozen global model was already unreliable: publication-balanced R2 = -0.075, MAE = 3.838 percentage points, and bias = -2.321 percentage points. The domain expert further reduced R2 by 0.462 and increased MAE by 18.2% (95% CI, 7.5% to 30.7%); its bias deteriorated to -3.187 percentage points. A negative R2 indicates that the subgroup predictions were inferior to a publication-balanced constant-mean reference on the independent future cohort.",
        bold_lead="Mixed Pb-Sn.",
    )

    add_table_caption(doc, "Table 3. Domain-expert MAE change across photovoltaic targets")
    table = doc.add_table(rows=1, cols=5)
    headers = ["Domain", "Target", "MAE change", "Paired 95% CI", "Interpretation"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
    expert = paired.loc[
        paired["method"].eq("Domain expert")
        & paired["evaluation"].eq("Publication-balanced")
    ]
    for domain in ["Sn-only", "Mixed Pb-Sn"]:
        for target in ["PCE", "Voc", "Jsc", "FF"]:
            row = expert.loc[expert["domain"].eq(domain) & expert["target"].eq(target)].iloc[0]
            if row.MAE_change_percent_CI_low > 0:
                interpretation = "Significant worsening"
            elif row.MAE_change_percent_CI_high < 0:
                interpretation = "Significant improvement"
            else:
                interpretation = "Inconclusive"
            values = [
                domain,
                target,
                f"{row.MAE_change_percent:+.1f}%",
                f"[{row.MAE_change_percent_CI_low:+.1f}, {row.MAE_change_percent_CI_high:+.1f}]",
                interpretation,
            ]
            cells = table.add_row().cells
            for cell, text in zip(cells, values):
                cell.text = text
    style_table(table, [1550, 900, 1500, 2100, 3310], numeric_columns={1, 2, 3})

    doc.add_heading("Upper-tail and OOD behavior", level=1)
    sn_upper = upper.loc[
        upper["domain"].eq("Sn-only")
        & upper["method"].eq("Frozen global")
        & upper["subset"].str.startswith("Historical")
    ].iloc[0]
    mixed_upper = upper.loc[
        upper["domain"].eq("Mixed Pb-Sn")
        & upper["method"].eq("Frozen global")
        & upper["subset"].str.startswith("Historical")
    ].iloc[0]
    add_body(
        doc,
        f"Upper-tail thresholds were fixed from the historical publication-balanced 75th percentile, not from future outcomes. For Sn-only devices the threshold was {sn_upper.threshold_PCE:.3f}% PCE; the future upper tail contained {int(sn_upper.records)} records from {int(sn_upper.DOI_groups)} DOI groups and retained a global bias of {sn_upper.bias:.3f} percentage points. For mixed Pb-Sn devices the threshold was {mixed_upper.threshold_PCE:.3f}% PCE; {int(mixed_upper.records)} records from {int(mixed_upper.DOI_groups)} DOI groups showed a global bias of {mixed_upper.bias:.3f} percentage points. The domain experts made both upper-tail biases more negative.",
    )
    mixed_unseen_global = robustness.loc[
        robustness["stratum"].eq("Formula unseen historically")
        & robustness["domain"].eq("Mixed Pb-Sn")
        & robustness["target"].eq("PCE")
        & robustness["method"].eq("Frozen global")
    ].iloc[0]
    mixed_unseen_expert = robustness.loc[
        robustness["stratum"].eq("Formula unseen historically")
        & robustness["domain"].eq("Mixed Pb-Sn")
        & robustness["target"].eq("PCE")
        & robustness["method"].eq("Domain expert")
    ].iloc[0]
    add_body(
        doc,
        f"The hardest mixed Pb-Sn stratum comprised historically unseen formulas ({int(mixed_unseen_global.records)} records, {int(mixed_unseen_global.DOI_groups)} DOI groups). Global PCE MAE was {mixed_unseen_global.MAE:.3f} with bias {mixed_unseen_global.bias:.3f}; the subgroup expert increased MAE to {mixed_unseen_expert.MAE:.3f} and bias magnitude to {mixed_unseen_expert.bias:.3f}. This pattern supports a missing-support explanation rather than a simple domain-wide intercept error.",
    )
    add_note(
        doc,
        "HIGH-EFFICIENCY LIMIT",
        "No Sn-only future device reached 20% PCE, and the mixed Pb-Sn >=20% subset contained only two DOI groups. "
        "Its global mean bias (-8.286 percentage points) is scientifically important but strictly descriptive; it is not used for inferential model selection.",
    )

    doc.add_heading("Decision and implications", level=1)
    add_body(
        doc,
        "The correct final policy is not to deploy a separate Sn expert with the present database. The one-standard-error rule treated small historical gains as indistinguishable from sampling noise and shrank all calibrators to identity and all expert mixture weights to zero. The independent test validated that conservative choice: fully domain-specific PCE models were significantly worse in both domains.",
    )
    add_body(
        doc,
        "This does not establish that subgroup modelling is intrinsically ineffective. It establishes that 83 historical Sn-only DOI groups and 57 historical mixed Pb-Sn DOI groups do not provide enough stable coverage of rapidly changing compositions and processes to justify a separate high-dimensional Random Forest. The global model should therefore remain the point predictor, while Sn-only and mixed Pb-Sn flags should trigger reliability warnings, OOD reporting, and wider interpretation rather than automatic recalibration.",
    )
    add_bullet(doc, "Primary model: retain the frozen full 1/n_DOI global Random Forest.")
    add_bullet(doc, "Reliability policy: label mixed Pb-Sn predictions as limited-support, particularly for unseen formulas and upper-tail PCE.")
    add_bullet(doc, "Prospective update rule: reconsider a subgroup expert only after materially larger independent DOI support and a locked chronological evaluation protocol.")
    add_bullet(doc, "Manuscript framing: report the negative adaptation result as evidence that average domain labels cannot substitute for detailed composition/process support.")

    doc.add_heading("Manuscript-ready Methods text", level=1)
    doc.add_heading("Subgroup calibration and mixture-of-experts", level=2)
    add_body(
        doc,
        "We evaluated whether composition-specific adaptation improved transfer to Sn-containing perovskite solar cells. Two absorber domains were defined before modelling from parsed B-site composition: Sn-only devices contained Sn but no Pb, whereas mixed Pb-Sn devices contained both Pb and Sn. The reference predictor was the previously selected multi-output Random Forest trained with full inverse-publication weighting, such that each DOI contributed equal total training weight. DOI, publication year, and photovoltaic outcomes were excluded from the feature matrix.",
    )
    add_body(
        doc,
        "Model development was restricted to records published through 2018. Historical out-of-fold predictions were generated using the frozen five DOI-group folds, with preprocessing and model fitting repeated within each training partition. Three adaptation strategies were considered for PCE, Voc, Jsc, and FF. First, subgroup-specific point calibration compared identity, bias-only, and affine residual corrections with ridge penalties of 0, 0.001, 0.01, 0.1, 1, and 10. Second, subgroup experts were trained only on the corresponding absorber domain; candidate Random Forest configurations varied categorical/token frequency thresholds and minimum leaf size. Third, global and expert predictions were combined as p_mix = (1-alpha)p_global + alpha p_expert, with alpha evaluated from 0 to 1 in increments of 0.05.",
    )
    add_body(
        doc,
        "All calibration choices, expert hyperparameters, mixture weights, and final method choices were based exclusively on historical DOI-disjoint predictions. We applied a one-standard-error rule at the DOI level, favoring identity calibration, lower expert complexity, and smaller alpha when performance was statistically indistinguishable from the minimum cross-validated error. The selected models were then refitted using all eligible historical records and evaluated once on the independent 2019-2021 chronological cohort. Test outcomes were not used for tuning.",
    )
    add_body(
        doc,
        "Performance was evaluated at the device level and with publication-balanced weights that assigned equal total evaluation weight to each test DOI. R2, MAE, RMSE, and mean prediction bias were reported. Differences from the frozen global model were calculated on identical records using a paired DOI-cluster bootstrap with 1,000 replicates. Subsets with fewer than 20 DOI groups were treated as descriptive. Upper-tail PCE thresholds were fixed as the historical publication-balanced 75th percentile within each domain."
    )

    doc.add_heading("Manuscript-ready Results and Discussion", level=1)
    doc.add_heading("Subgroup-specific adaptation did not rescue Sn-domain transfer", level=2)
    add_body(
        doc,
        "The composition audit identified two scientifically important but sparsely supported Sn-containing domains. The historical development cohort included 389 Sn-only devices from 83 publications and 274 mixed Pb-Sn devices from 57 publications. Their independent chronological test cohorts contained 155 devices from 30 publications and 98 devices from 25 publications, respectively. We therefore tested deliberately regularized adaptation strategies rather than fitting an unrestricted gate on the future data.",
    )
    add_body(
        doc,
        "Historical publication-disjoint selection provided no stable evidence for subgroup calibration or expert substitution. For every domain-target combination, identity calibration lay within one standard error of the minimum candidate error. Likewise, the convex expert weight was shrunk to alpha = 0 for PCE, Voc, Jsc, and FF, and the development-selected policy retained the frozen global model. These results show that small apparent cross-validation gains were insufficiently robust relative to between-publication variability.",
    )
    add_body(
        doc,
        "The untouched chronological test supported this conservative decision. For Sn-only PCE, the global model achieved a publication-balanced MAE of 1.985 percentage points and R2 of 0.230, whereas the domain expert increased MAE to 2.271 percentage points. The paired increase was 14.4% (95% CI, 8.0-22.1%), accompanied by a reduction in R2 of 0.169 (95% CI, 0.093-0.249). For mixed Pb-Sn PCE, global performance was already poor (MAE, 3.838 percentage points; R2, -0.075; bias, -2.321 percentage points). The subgroup expert further increased MAE to 4.537 percentage points, corresponding to an 18.2% deterioration (95% CI, 7.5-30.7%), and made the underprediction bias more severe (-3.187 percentage points)."
    )
    add_body(
        doc,
        "The failure was not limited to PCE. Sn-only experts significantly worsened Jsc and FF MAE, while mixed Pb-Sn experts significantly worsened FF and produced no conclusive gain for Voc or Jsc. Although the mixed Pb-Sn expert reduced Jsc MAE by 5.6% on average, its paired 95% confidence interval ranged from an 18.7% improvement to a 4.6% deterioration. This isolated, uncertain change does not justify a target-specific deployment rule.",
    )
    add_body(
        doc,
        "Error concentration in emerging compositions explains why a coarse domain expert was ineffective. Among mixed Pb-Sn formulas absent from the historical corpus, global PCE MAE was 5.227 percentage points with a bias of -3.421 percentage points; restricting the model to historical mixed Pb-Sn records increased MAE to 6.593 percentage points and bias magnitude to -5.360 percentage points. Similarly, both domains retained pronounced underprediction in their historically defined upper-performance quartiles. The sparse expert therefore removed useful cross-domain information without supplying enough within-domain coverage of new formulas and processes.",
    )
    add_body(
        doc,
        "These findings refine the interpretation of domain-specific modelling. A subgroup label can identify elevated risk, but it does not guarantee that a separately trained model will be better calibrated. With only tens of independent publications, publication-level heterogeneity dominates the modest benefits of specialization. We therefore retain the full DOI-balanced global model and use Sn-only and mixed Pb-Sn membership as reliability flags coupled to OOD diagnostics. A future subgroup expert should be reconsidered prospectively only after additional independent publications expand coverage, especially in mixed-composition and high-efficiency regimes."
    )

    doc.add_heading("Reproducibility and integrity", level=1)
    add_body(
        doc,
        f"The analysis generated {manifest['prediction_rows']:,} future prediction rows with zero duplicate keys. Historical and future DOI overlap was zero. The frozen global predictions matched the archived chronological predictions to within 3.55e-15, and independently recalculated summary metrics agreed to within 1.78e-15. The verification suite passed {verification['checks_passed']} of {verification['checks_total']} checks. Random seeds, model configurations, input hashes, selection tables, predictions, and all 1,000-replicate bootstrap summaries are included in the accompanying reproducibility package.",
    )

    add_body(doc, "The accompanying reproducibility package contains the complete analysis script, independent verifier, compressed row-level future predictions, selection tables, performance and paired-comparison tables, upper-tail/OOD summaries, the run manifest, and Figure 9 in PNG, PDF, and SVG formats.")

    return doc


def main() -> None:
    document = build_document()
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
